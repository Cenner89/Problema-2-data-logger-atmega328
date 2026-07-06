from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RELATORIO_TECNICO_FORMATADO.docx"


def font(run, size=10, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def columns(section, count):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if cols.getparent() is None:
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "720")


def p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10, bold=False, italic=False, first=False):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_after = Pt(3)
    if first:
        par.paragraph_format.first_line_indent = Cm(0.45)
    run = par.add_run(text)
    font(run, size=size, bold=bold, italic=italic)
    return par


def heading(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text)
    font(run, size=10, bold=True)
    return par


def subheading(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_before = Pt(5)
    par.paragraph_format.space_after = Pt(2)
    run = par.add_run(text)
    font(run, size=10, bold=True)
    return par


def code_line(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)
    return par


def code_block(doc, lines):
    for line in lines:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.left_indent = Cm(0.25)
        par.paragraph_format.space_after = Pt(0)
        run = par.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(7)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(1.7)
sec.right_margin = Cm(1.7)

for style_name in ["Normal"]:
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(10)

p(doc, "©2026 IEEE", align=WD_ALIGN_PARAGRAPH.LEFT, size=8)

title = p(
    doc,
    "Projeto de um Data-Logger Inteligente para\nMonitoramento Ambiental Remoto Utilizando o\nMicrocontrolador ATmega328P",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=16,
    bold=True,
)
title.paragraph_format.space_after = Pt(10)

authors = doc.add_table(rows=1, cols=3)
authors.alignment = WD_TABLE_ALIGNMENT.CENTER
authors.autofit = True
author_data = [
    ("Brunna Mota", "Departamento de Engenharia Elétrica e\nde Computação\nUniversidade Federal da Bahia\nSalvador, Brasil\nbrunna.mota@ufba.br"),
    ("Matheus Correia", "Departamento de Engenharia Elétrica e\nde Computação\nUniversidade Federal da Bahia\nSalvador, Brasil\nmatheushcc@ufba.br"),
    ("Werner Guaraná", "Departamento de Engenharia Elétrica e\nde Computação\nUniversidade Federal da Bahia\nSalvador, Brasil\nwguarana@gmail.com"),
]
for cell, (name, info) in zip(authors.rows[0].cells, author_data):
    cell.text = ""
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(name + "\n")
    font(r, 10, bold=False)
    r = par.add_run(info)
    font(r, 9)

p(doc)

abstract = doc.add_paragraph()
abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = abstract.add_run("Resumo — ")
font(r, 10, bold=True)
r = abstract.add_run(
    "Este trabalho apresenta o desenvolvimento de uma base de firmware para um data-logger ambiental usando o microcontrolador ATmega328P. "
    "O sistema foi pensado para coletar dados de sensores analógicos e digitais, associar cada medição a data e hora usando um RTC externo, "
    "aplicar um filtro digital passa-baixas nas leituras analógicas e transmitir os registros por comunicação serial. Também foi incluída uma estrutura inicial para comunicação sem fio, usando a mesma lógica de envio dos dados. "
    "Além disso, foi implementado um armazenamento simples em EEPROM interna, com funcionamento circular, para manter os registros mais recentes. "
    "A base foi validada por compilação com MPLAB XC8 e por um teste simples da lógica do filtro. Algumas partes ainda precisam de validação em bancada, principalmente sensores reais, RTC físico e módulo sem fio."
)
font(r, 10)

kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = kw.add_run("Palavras-chave — ")
font(r, 10, bold=True, italic=True)
r = kw.add_run("ATmega328P, Data-Logger, Sistemas Embarcados, Monitoramento Ambiental, RTC, EEPROM, Filtragem Digital.")
font(r, 10, italic=True)

doc.add_section(WD_SECTION.CONTINUOUS)
columns(doc.sections[-1], 2)

heading(doc, "I. INTRODUÇÃO")
p(doc, "O monitoramento ambiental é importante em aplicações como agricultura, acompanhamento climático, preservação ambiental e análise de áreas remotas. Em muitos casos, a coleta desses dados ainda é feita manualmente, o que pode gerar atrasos, diferenças nos horários das medições e até perda de informação.", first=True)
p(doc, "A proposta deste trabalho foi montar a base de um sistema embarcado capaz de registrar variáveis ambientais automaticamente. Para isso, foi usado o ATmega328P como microcontrolador principal. O sistema faz leitura de sensores, processa os dados, associa data e hora às medições e transmite as informações em formato simples.", first=True)
p(doc, "O desenvolvimento foi feito em C, pois o problema pediu essa linguagem e porque ela facilita bastante a organização do firmware quando comparada com Assembly. Como o projeto envolve vários periféricos, como ADC, USART, I2C e EEPROM, separar o código em módulos foi uma decisão importante.", first=True)

heading(doc, "II. ARQUITETURA DO SISTEMA")
p(doc, "A arquitetura do sistema foi pensada de forma modular. O ATmega328P fica responsável por fazer a leitura dos sensores, aplicar o filtro, buscar data e hora no RTC, armazenar os dados na EEPROM e enviar os registros por comunicação serial ou sem fio.", first=True)
p(doc, "Foram considerados quatro sensores ambientais. Os sensores de temperatura e luminosidade são tratados como analógicos e ligados ao ADC. Já os sensores de umidade e chuva são tratados como digitais, usando entradas GPIO do microcontrolador.", first=True)
p(doc, "Depois da leitura, os valores analógicos passam por um filtro digital passa-baixas. Em seguida, o registro recebe o timestamp vindo do RTC DS3231. Por fim, os dados são armazenados na EEPROM e enviados no formato CSV.", first=True)
heading(doc, "III. Implementação do Sistema")
p(doc, "A implementação foi organizada em arquivos separados para deixar o projeto mais fácil de entender. No começo até daria para colocar tudo em um arquivo só, mas isso deixaria o main.c muito grande e confuso. Como o projeto tem sensor, filtro, RTC, comunicação e armazenamento, separar por função acabou sendo mais organizado.", first=True)
p(doc, "Na prática, o main.c fica quase só chamando o data-logger. Quem coordena o sistema é o módulo datalogger, mostrado no trecho abaixo. Ele pega o horário, lê os sensores, salva o registro e envia os dados.", first=True)
code_block(doc, [
    "void datalogger_run_once(void)",
    "{",
    "    log_record_t record;",
    "",
    "    (void)rtc_get_time(&record.timestamp);",
    "    sensors_read(&record.sensors);",
    "    (void)storage_save_record(&record);",
    "    comm_send_record(&record);",
    "}",
])
p(doc, "Essa divisão não foi feita para deixar o código mais bonito apenas. Ela ajuda bastante na hora de mexer no projeto, porque se der problema na comunicação, por exemplo, não é preciso procurar no meio do código do ADC ou do filtro.", first=True)

subheading(doc, "A. Aquisição de Dados")
p(doc, "A aquisição dos dados foi dividida entre sensores analógicos e digitais. Os sensores analógicos usam o ADC do ATmega328P. Nesta proposta, o sensor de temperatura foi colocado no canal ADC0 e o sensor de luminosidade, baseado em LDR, no canal ADC1.", first=True)
p(doc, "Os sensores digitais de umidade e chuva foram configurados como entradas digitais. Para evitar entradas flutuantes durante os testes, foram usados pull-ups internos. Como ainda não foram definidos os modelos exatos dos sensores, o firmware trabalha inicialmente com valores brutos.", first=True)

subheading(doc, "B. Filtragem Digital")
p(doc, "Os sensores analógicos podem apresentar oscilações por ruído elétrico ou variações pequenas do próprio circuito. Para suavizar isso, foi implementado um filtro digital passa-baixas do tipo IIR simples.", first=True)
code_line(doc, "y = y + ((x - y) >> FILTER_SHIFT)")
p(doc, "No código, essa ideia aparece de forma bem direta:", first=True)
code_block(doc, [
    "delta = (int16_t)sample - (int16_t)filter->value;",
    "filter->value = (uint16_t)((int16_t)filter->value +",
    "                           (delta >> FILTER_SHIFT));",
])
p(doc, "Essa escolha evita o uso de ponto flutuante e usa apenas operações inteiras, o que combina melhor com o ATmega328P. No projeto, FILTER_SHIFT foi definido como 3, equivalente a uma suavização aproximada de 1/8 da diferença entre a nova leitura e o valor filtrado anterior.", first=True)

subheading(doc, "C. Gerenciamento de Data e Hora")
p(doc, "Para registrar o instante de cada medição, foi adotado o RTC DS3231. Ele se comunica com o ATmega328P pelo barramento I2C/TWI, usando os pinos PC4 para SDA e PC5 para SCL.", first=True)
p(doc, "O firmware lê os registradores do DS3231, converte os valores de BCD para decimal e monta uma estrutura com ano, mês, dia, hora, minuto e segundo. Caso o RTC não responda, o sistema usa um horário padrão, apenas para permitir que o restante do firmware continue funcionando.", first=True)

subheading(doc, "D. Armazenamento dos Dados")
p(doc, "Os registros são armazenados na EEPROM interna do ATmega328P. Foi usada uma estratégia circular: quando o número máximo de registros é atingido, os dados novos passam a sobrescrever os mais antigos.", first=True)
p(doc, "Essa solução não é a mais completa para um produto real, porque a EEPROM tem pouco espaço e limite de ciclos de escrita. Mesmo assim, ela atende bem à proposta do trabalho, pois demonstra o funcionamento básico de um data-logger sem precisar de memória externa.", first=True)

subheading(doc, "E. Comunicação")
p(doc, "A comunicação foi organizada em duas camadas. O módulo usart cuida diretamente da transmissão serial, enquanto o módulo comm monta os registros no formato definido.", first=True)
p(doc, "O formato escolhido foi CSV, porque é fácil de visualizar em um terminal serial e também pode ser importado em planilhas. O cabeçalho enviado pelo sistema é:", first=True)
code_line(doc, "timestamp,temp_raw,temp_filt,light_raw,light_filt,humidity,rain,mode")
p(doc, "Um trecho simplificado da montagem do registro é:", first=True)
code_block(doc, [
    "comm_send_timestamp(record);",
    "usart_send_char(',');",
    "usart_send_uint16(record->sensors.temperature_raw);",
    "usart_send_char(',');",
    "usart_send_uint16(record->sensors.temperature_filtered);",
])
p(doc, "Para a comunicação sem fio, a ideia foi manter o mesmo formato de dados e usar um módulo compatível com comunicação serial, como um módulo Bluetooth. Assim, o firmware muda pouco entre a versão com cabo e a versão sem fio.", first=True)

heading(doc, "Conclusão")
p(doc, "O projeto resultou em uma base funcional para um data-logger ambiental com ATmega328P. Foram implementados os principais blocos pedidos no problema: leitura de sensores, filtro digital, RTC, armazenamento, comunicação e seleção da interface.", first=True)
p(doc, "O firmware compila com MPLAB XC8 e gera o arquivo HEX. Também foi criado um teste simples para validar a lógica do filtro passa-baixas. Apesar disso, ainda faltam testes físicos com sensores reais, RTC DS3231 conectado e módulo sem fio. Portanto, o trabalho está fechado como uma base de projeto organizada e pronta para validação em bancada.", first=True)

refs = [
    "[1] Microchip Technology Inc., ATmega328P Datasheet.",
    "[2] Microchip Technology Inc., MPLAB XC8 C Compiler User's Guide.",
    "[3] Analog Devices / Maxim Integrated, DS3231 Extremely Accurate I2C-Integrated RTC.",
    "[4] M. A. Mazidi, S. Naimi, S. Naimi, The AVR Microcontroller and Embedded Systems Using Assembly and C.",
    "[5] Material da disciplina ENGC50 - Sistemas Microprocessados, 2026.1.",
]
for ref in refs:
    p(doc, ref, size=8)

doc.save(OUT)
print(OUT)
